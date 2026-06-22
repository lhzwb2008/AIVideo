"""生成股市论坛手动发文包：与视频同目录的同名文件夹 post.md + images/ + cover.jpg + cover_landscape.jpg。"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
from pathlib import Path

from paths import ROOT
from research import extract_json, load_env
from text_client import chat_complete

DISCLAIMER = ""

# 财富号/雪球：弱化标题党、连板炒作等表述
_FORUM_TITLE_REPLACEMENTS: tuple[tuple[str, str], ...] = (
    (r"五连板", "连续上涨"),
    (r"为什么能连续上涨", "近期为何表现强势"),
    (r"单日暴跌", "单日大幅回调"),
    (r"为啥", "为何"),
)

# 正文：互动引流、荐股式结尾（东方财富 2.1）
_CTA_PATTERNS: tuple[str, ...] = (
    r"觉得有用就?点赞收藏[，,、]?关注我们[^。！？?]*[。！？?]?\s*$",
    r"点赞收藏[，,、]?关注我们[^。！？?]*[。！？?]?\s*$",
    r"关注我们[，,、]?每天[^。！？?]*[。！？?]?\s*$",
    r"觉得有用点个关注!?\s*$",
    r"评论区聊聊[^。！？?]*[。！？?]?\s*$",
    r"那你觉得[^。！？?]*评论区[^。！？?]*[。！？?]?\s*$",
    r"那么问题来了[：:]?[^。！？?]*你看好[^。！？?]*[。！？?]?\s*$",
    r"那问题来了[：:]?[^。！？?]*(真行情|凑热闹|五连板)[^。！？?]*[。！？?]?\s*$",
    r"你觉得[^。！？?]*(涌向|流向)哪个(方向|板块)[？?]\s*$",
    r"你觉得[^。！？?]*[？?]\s*$",
    r"那你觉得[^。！？?]*[？?]\s*$",
    r"问题来了[：:]?[^。！？?]*[？?]\s*$",
)

# 正文：点名个股 + 极端涨跌幅（东方财富 2.2 / 2.4）
_FORUM_NARRATION_REPLACEMENTS: tuple[tuple[str, str], ...] = (
    (
        r"有只煤炭股一口气连续五个交易日涨停，一周累计涨了61%！这是啥概念？就好比一家网红奶茶店连续五天爆单到打烊都排不上号。它叫华电能源，这周的明星。",
        "部分煤炭股却明显强于指数，个别标的短期涨幅较大。同一市场里，个股走势分化很大。就好比一条美食街上，有的店门庭若市，有的店冷冷清清。",
    ),
    (
        r"一边是26只股票涨超30%，地产股香江控股也走出五连板；另一边呢，有76只股票直接跌超20%，其中朗信电气一周跌了45%，几乎腰斩。",
        "一边是个别标的涨幅靠前，另一边也有大量个股明显回调，跌幅超过两成的不在少数。",
    ),
    (
        r"这周542家公司被调研，神工股份最火，55家机构同一周排队上门看账本。机构调研就好比一群投资人专门上门翻这家店的账本，看看值不值得长期关注。",
        "这周五百多家公司接受机构调研，部分半导体、制造类公司关注度较高。机构调研就好比投资人专门翻阅企业资料，评估长期价值。需注意的是，短期涨幅较大的标的往往波动也大，追高风险不容忽视。",
    ),
    (r"中芯国际跌了近9%", "部分龙头芯片股跌幅明显"),
    (r"公共事业", "公用事业"),
    (r"EBITA", "EBITDA"),
    (
        r"雷神科技盘中直接冲到30%涨停",
        "部分AI PC概念股盘中涨幅明显",
    ),
    (
        r"今年累计涨了793%的8倍大牛股利通电子，却来了个一字跌停",
        "部分前期涨幅较大的高位标的，同日也出现明显回调",
    ),
)


def forum_dir_for_video(video_path: Path) -> Path:
    """与 mp4 同级、同名文件夹，如 output/20260531_193024/"""
    return video_path.parent / video_path.stem


def _load_script(path: Path) -> dict:
    data = json.loads(path.read_text(encoding="utf-8"))
    script = data.get("script") or data
    if not isinstance(script, dict):
        raise ValueError(f"无效脚本: {path}")
    return script


def _load_script_meta(path: Path) -> tuple[dict, dict]:
    data = json.loads(path.read_text(encoding="utf-8"))
    script = data.get("script") or data
    if not isinstance(script, dict):
        raise ValueError(f"无效脚本: {path}")
    return script, data if isinstance(data, dict) else {}


def _sanitize_forum_title(title: str) -> str:
    t = title.strip()
    for pat, repl in _FORUM_TITLE_REPLACEMENTS:
        t = re.sub(pat, repl, t)
    return t.strip()


def _sanitize_forum_narration(text: str) -> str:
    t = text.strip()
    for pat, repl in _FORUM_NARRATION_REPLACEMENTS:
        t = re.sub(pat, repl, t)
    return t.strip()


def _strip_cta(text: str) -> str:
    t = text.strip()
    changed = True
    while changed and t:
        changed = False
        for pat in _CTA_PATTERNS:
            new_t = re.sub(pat, "", t).strip()
            if new_t != t:
                t = new_t
                changed = True
                break
    return t


def _prepare_forum_narration(text: str) -> str:
    return _strip_cta(_sanitize_forum_narration(text))


def _label_covered(label: str, narration: str) -> bool:
    label = label.strip()
    if not label or label in narration:
        return True
    core = label
    for sep in ("=", "＝", "：", ":"):
        if sep in core:
            core = core.split(sep, 1)[0].strip()
            break
    chunks = re.findall(r"[\u4e00-\u9fff]{3,}", core)
    if not chunks:
        return core in narration
    hits = sum(1 for chunk in chunks if chunk in narration)
    return hits >= max(1, len(chunks) - 1)


def _label_to_sentence(label: str) -> str:
    label = _prepare_forum_narration(label.strip())
    if not label:
        return ""
    if label.endswith(("。", "！", "？")):
        return label
    for sep in ("=", "＝", "：", ":"):
        if sep in label:
            key, val = label.split(sep, 1)
            key, val = key.strip(), val.strip()
            if key and val and len(val) >= 2:
                return f"{key}方面，大致对应{val.rstrip('。')}。"
    return ""


def _expand_from_labels(labels: list[str], narration: str) -> str:
    sentences: list[str] = []
    for lb in labels:
        if _label_covered(lb, narration):
            continue
        sent = _label_to_sentence(lb)
        if sent and sent not in narration and sent not in sentences:
            sentences.append(sent)
    if not sentences:
        return ""
    if len(sentences) == 1:
        return sentences[0]
    return "补充几个要点：" + "".join(sentences)


def _expand_forum_section(slide: dict) -> str:
    narration = _prepare_forum_narration(str(slide.get("narration") or ""))
    labels = [str(x).strip() for x in (slide.get("on_image_text") or []) if str(x).strip()]
    parts: list[str] = []

    if narration:
        parts.append(narration)

    extra = _expand_from_labels(labels, "\n".join(parts))
    if extra:
        parts.append(extra)

    concept = _prepare_forum_narration(str(slide.get("concept") or ""))
    headline = str(slide.get("headline") or "").strip()
    if (
        concept
        and concept != headline
        and len(concept) >= 8
        and not narration.startswith(concept[: min(6, len(concept))])
    ):
        from forum_editor_fill import concept_redundant

        if not concept_redundant(concept, "\n".join(parts)):
            parts.append(concept.rstrip("。") + "。")

    return "\n\n".join(p for p in parts if p.strip())


def _extract_video_frames(video: Path, dest_dir: Path, count: int) -> list[Path]:
    dest_dir.mkdir(parents=True, exist_ok=True)
    dur_s = 20.0 * count
    try:
        out = subprocess.check_output(
            [
                "ffprobe",
                "-v",
                "error",
                "-show_entries",
                "format=duration",
                "-of",
                "default=noprint_wrappers=1:nokey=1",
                str(video),
            ],
            text=True,
        )
        dur_s = float(out.strip())
    except (subprocess.CalledProcessError, FileNotFoundError, ValueError):
        pass

    paths: list[Path] = []
    for i in range(count):
        t = dur_s * (i + 0.5) / max(count, 1)
        out = dest_dir / f"{i + 1:02d}.jpg"
        subprocess.run(
            [
                "ffmpeg",
                "-y",
                "-ss",
                f"{t:.2f}",
                "-i",
                str(video),
                "-frames:v",
                "1",
                "-q:v",
                "2",
                str(out),
            ],
            capture_output=True,
            check=False,
        )
        if out.is_file():
            paths.append(out)
    return paths


def _copy_slide_images(script: dict, dest_dir: Path, video: Path, n_slides: int) -> list[Path]:
    dest_dir.mkdir(parents=True, exist_ok=True)
    slides = script.get("slides") or []
    copied: list[Path] = []
    for i, slide in enumerate(slides[:n_slides], start=1):
        rel = slide.get("image_path") or slide.get("cover_image")
        if rel:
            src = ROOT / rel
            if src.is_file():
                dst = dest_dir / f"{i:02d}.jpg"
                shutil.copy2(src, dst)
                copied.append(dst)
    if len(copied) >= max(1, min(n_slides, 1)):
        return copied[:n_slides] if len(copied) >= n_slides else copied
    cover = script.get("cover_image")
    if cover:
        src = ROOT / cover
        if src.is_file():
            dst = dest_dir / "01.jpg"
            shutil.copy2(src, dst)
            copied = [dst]
    if len(copied) >= n_slides:
        return copied[:n_slides]
    if video.is_file():
        extracted = _extract_video_frames(video, dest_dir, n_slides)
        if extracted:
            return extracted
    thumb = ROOT / "logs/youtube_thumbs" / f"{video.stem}_frame0.jpg"
    if thumb.is_file():
        dst = dest_dir / "01.jpg"
        shutil.copy2(thumb, dst)
        return [dst]
    return copied


def _save_cover_jpg(src: Path, cover_dst: Path) -> bool:
    if src.suffix.lower() in {".jpg", ".jpeg"}:
        shutil.copy2(src, cover_dst)
        return cover_dst.is_file()
    try:
        proc = subprocess.run(
            ["ffmpeg", "-y", "-i", str(src), "-q:v", "2", str(cover_dst)],
            capture_output=True,
            check=False,
            timeout=120,
        )
    except (OSError, subprocess.TimeoutExpired):
        return False
    return proc.returncode == 0 and cover_dst.is_file()


def _write_cover(script_path: Path, video: Path, out_dir: Path) -> tuple[Path | None, Path | None]:
    """论坛竖封面 + 源图路径（供横封面裁剪）。竖图 = 视频开场，不用正文配图。"""
    from publish_resolve import resolve_cover_image

    src = resolve_cover_image(script_path, video)
    if not src or not src.is_file():
        return None, None
    cover_dst = out_dir / "cover.jpg"
    if not _save_cover_jpg(src, cover_dst):
        return None, src
    return cover_dst, src


def _quote_block(text: str) -> str:
    return f"```\n{text.rstrip()}\n```\n\n"


def _build_readme_publish_sections(script: dict) -> str:
    """视频平台通用发布文案（抖音/小红书/视频号等同），写入 README 便于归档后一键复制。"""
    from douyin_caption import build_sau_fields

    dy = build_sau_fields(script)
    dy_tags = (dy.get("tags") or "").strip()
    dy_hashtags = " ".join(f"#{t.strip()}" for t in dy_tags.split(",") if t.strip())
    dy_desc = (dy.get("desc") or "").strip()
    if dy_hashtags:
        dy_desc = f"{dy_desc}\n\n{dy_hashtags}".strip()

    return (
        "## 发布文案（可直接复制）\n\n"
        "### 抖音\n\n"
        f"**标题**\n\n{_quote_block(dy['title'])}"
        f"**标签**（逗号分隔，发布时选话题）\n\n{_quote_block(dy_tags)}"
        f"**简介 + 话题**（整段复制）\n\n{_quote_block(dy_desc)}"
    )


FORUM_ARTICLE_SYSTEM = """你是「AI财知道」财经图文编辑。

任务：根据短视频分镜脚本，写一篇适合雪球/东方财富发布的图文长文。读者会边看与视频相同的漫画分镜配图边读正文，因此叙事必须与视频一致，写法则从口语改为文章。

核心原则（折中，必须遵守）：
1. **叙事骨架 = 视频分镜**：严格按 slides 顺序，第 i 页对应第 i 节；不得打乱顺序，不得改成另一套独立研究框架或换话题。
2. **写法改书面、逻辑不变**：把口播 narration 改写成财经科普/评论文章（完整句、因果与转折、2-3 段自然段）；去掉互动引导、点赞关注、评论区话术；口播式夸张比喻改为克制表述，可保留少量有助于理解的类比。
3. **呼应画面**：每节正文要解释该页漫画在讲什么，自然融入 on_image_text 里的要点，让读者看图能读懂。
4. **可补充细节**：在【深读材料】【文章来源】中选取与当页叙事相关的数字、背景、引语补充进来；禁止引入与视频主线无关的新章节或跑题。
5. 全文 1200-2000 字，信息密度高于口播总字数，但不必写成券商深度报告体例。
6. 合规：不荐股、不喊单、不预测买卖点；不用“必涨/抄底/梭哈/能不能买”等表述；审慎用语（可能、尚需观察、取决于）。
7. 输出 JSON，不要 markdown。
"""

FORUM_ARTICLE_USER = """请生成图文长文。必须输出 **恰好 {slide_count} 个 sections**，与下方【视频分镜】逐页一一对应（第 1 节 = 第 1 页，以此类推）。

【标题】
{title}

【栏目】
{category}

【冷开场】
{cold_open}

【文章来源】（仅补充与视频同主线的事实，勿改叙事顺序）
{article_json}

【深读材料】（仅补充与当页相关的数字/背景/引语）
{details_json}

【视频分镜：每页 headline / narration / on_image_text 即该节叙事骨架】
{script_json}

输出格式：
{{
  "title": "适合雪球/东方财富的标题，稳健清晰，不标题党",
  "sections": [
    {{
      "headline": "与小标题相近，可略书面化，勿偏离该页主题",
      "body": "2-3段正文，用\\n\\n分段；书面语；逻辑跟该页 narration 一致"
    }}
  ]
}}

sections 数组长度必须等于 {slide_count}。
"""


def _script_digest(script: dict) -> dict:
    slides = []
    for i, slide in enumerate(script.get("slides") or [], start=1):
        if not isinstance(slide, dict):
            continue
        slides.append(
            {
                "page": i,
                "chapter_title": slide.get("chapter_title"),
                "headline": slide.get("headline"),
                "concept": slide.get("concept"),
                "narration": slide.get("narration"),
                "on_image_text": slide.get("on_image_text"),
                "lead_in": slide.get("lead_in"),
            }
        )
    return {
        "title": script.get("title"),
        "category": script.get("category"),
        "keyword": script.get("keyword"),
        "cold_open": script.get("cold_open"),
        "hashtags": script.get("hashtags"),
        "slide_count": len(slides),
        "slides": slides,
    }


def _coerce_detail_list(value) -> list:
    """深读字段应为 list；模型偶发 dict 时先转成 list 再切片。"""
    if isinstance(value, list):
        return value
    if isinstance(value, dict):
        return [
            f"{k}：{v}" if str(v).strip() else str(k)
            for k, v in value.items()
            if str(k).strip()
        ]
    if value is None:
        return []
    text = str(value).strip()
    return [text] if text else []


def _details_digest(details: dict | None) -> dict:
    if not isinstance(details, dict):
        return {}
    return {
        "outline": _coerce_detail_list(details.get("outline"))[:16],
        "all_numbers": _coerce_detail_list(details.get("all_numbers"))[:24],
        "all_quotes": _coerce_detail_list(details.get("all_quotes"))[:12],
        "people": _coerce_detail_list(details.get("people"))[:12],
        "companies_or_institutions": _coerce_detail_list(details.get("companies_or_institutions"))[:16],
        "key_terms": _coerce_detail_list(details.get("key_terms"))[:16],
        "concrete_scenes": _coerce_detail_list(details.get("concrete_scenes"))[:10],
        "narrative_beats": _coerce_detail_list(details.get("narrative_beats"))[:12],
        "author_stance": details.get("author_stance"),
        "actual_opening": details.get("actual_opening"),
        "actual_ending": details.get("actual_ending"),
    }


def _article_digest(article: dict | None) -> dict:
    if not isinstance(article, dict):
        return {}
    return {
        "title": article.get("title"),
        "url": article.get("url"),
        "site": article.get("site"),
        "published_at": article.get("published_at"),
        "summary_zh": article.get("summary_zh"),
        "thesis": article.get("thesis"),
        "key_facts": (article.get("key_facts") or [])[:12],
        "narrative_arc": article.get("narrative_arc"),
    }


def _details_for_forum(script_meta: dict) -> dict | None:
    details = script_meta.get("research_details") or script_meta.get("details")
    if isinstance(details, dict) and details:
        return details
    fallback = ROOT / "logs" / "last_article_details.json"
    if fallback.is_file():
        try:
            data = json.loads(fallback.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else None
        except (OSError, json.JSONDecodeError):
            return None
    return None


def _normalize_forum_sections(rows: list[dict], slides: list[dict], image_paths: list[Path]) -> list[dict]:
    sections: list[dict] = []
    for i, row in enumerate(rows):
        if not isinstance(row, dict):
            continue
        headline = str(row.get("headline") or "").strip()
        body = str(row.get("body") or "").strip()
        if not headline and i < len(slides):
            headline = str(slides[i].get("headline") or "").strip()
        if not body:
            continue
        sections.append(
            {
                "headline": headline,
                "body": body,
                "image": image_paths[i] if i < len(image_paths) else None,
            }
        )
    return sections


def _plain_text_len(text: str) -> int:
    return len(re.sub(r"\s+", "", text or ""))


def _validate_professional_sections(
    title: str, sections: list[dict], *, slide_count: int
) -> None:
    del title
    body = "\n".join(str(s.get("body") or "") for s in sections)
    total_len = _plain_text_len(body)
    min_chars = int(os.environ.get("AIVIDEO_FORUM_ARTICLE_MIN_CHARS", "1000"))
    if total_len < min_chars:
        raise RuntimeError(f"长文字数不足：{total_len} 字，要求至少 {min_chars} 字")
    if slide_count and len(sections) != slide_count:
        raise RuntimeError(
            f"长文章节数须与视频分镜一致：{len(sections)} 节，要求 {slide_count} 节"
        )
    min_section = int(os.environ.get("AIVIDEO_FORUM_SECTION_MIN_CHARS", "120"))
    short_sections = [
        str(s.get("headline") or f"第{i + 1}节")
        for i, s in enumerate(sections)
        if _plain_text_len(str(s.get("body") or "")) < min_section
    ]
    if short_sections:
        raise RuntimeError(f"长文章节过短：{', '.join(short_sections[:3])}")


def _is_daily_recap_script(script: dict, article: dict | None = None) -> bool:
    plan = script.get("_topic_plan")
    if isinstance(plan, dict) and plan.get("script_mode") == "daily_recap":
        return True
    if isinstance(article, dict) and str(article.get("source_type") or "") == "cursor:astock_market":
        return True
    return False


def _generate_professional_forum_sections(
    script: dict,
    *,
    slides: list[dict],
    image_paths: list[Path],
    article: dict | None = None,
    details: dict | None = None,
) -> tuple[str, list[dict]]:
    load_env()
    slide_count = len(slides)
    base_user = FORUM_ARTICLE_USER.format(
        title=str(script.get("title") or ""),
        category=str(script.get("category") or ""),
        cold_open=str(script.get("cold_open") or ""),
        slide_count=slide_count,
        article_json=json.dumps(_article_digest(article), ensure_ascii=False, indent=2),
        details_json=json.dumps(_details_digest(details), ensure_ascii=False, indent=2),
        script_json=json.dumps(_script_digest(script), ensure_ascii=False, indent=2),
    )
    if _is_daily_recap_script(script, article):
        fixed = str(script.get("title") or "").strip()
        base_user += (
            "\n\n【长文类型：A股每日收盘报盘】\n"
            f"- 标题与视频一致：{fixed}\n"
            "- 4 节与视频分镜一一对应：①指数报盘 ②量能与涨跌家数 ③行业涨跌一览 ④一句话总结\n"
            "- 简单分析即可，禁止写成 MLCC/半导体/个股专题；行业名点到为止\n"
            "- 少用「加仓」「割肉」；资金用净流入/净流出\n"
        )
    plan = script.get("_topic_plan")
    if isinstance(plan, dict) and plan.get("offday_no_astock_recap"):
        base_user += (
            "\n\n【非交易日】正文禁止写 A 股指数收盘、4000点、成交额、涨跌家数；"
            "与视频一致，只写宏观/产业/国际主线。\n"
        )
    last_err: Exception | None = None
    user = base_user
    attempts = int(os.environ.get("AIVIDEO_FORUM_ARTICLE_RETRIES", "4"))
    for attempt in range(max(1, attempts)):
        raw = chat_complete(
            system=FORUM_ARTICLE_SYSTEM,
            user=user,
            max_tokens=int(os.environ.get("AIVIDEO_FORUM_ARTICLE_TOKENS", "6500")),
            response_format_json=True,
        )
        try:
            try:
                data = extract_json(raw)
            except ValueError:
                stripped = re.sub(
                    r"^```(?:json)?\s*|\s*```$", "", (raw or "").strip(), flags=re.MULTILINE
                )
                data = extract_json(stripped)
            title = _sanitize_forum_title(str(data.get("title") or script.get("title") or "未命名"))
            if _is_daily_recap_script(script, article):
                title = _sanitize_forum_title(str(script.get("title") or title))
            rows = data.get("sections") or data.get("paragraphs") or []
            if not rows and isinstance(data.get("body"), str) and data["body"].strip():
                rows = [{"headline": title, "body": data["body"].strip()}]
            sections = _normalize_forum_sections(rows, slides, image_paths)
            if not sections:
                preview = (raw or "").strip().replace("\n", " ")[:240]
                raise RuntimeError(
                    f"叙事长文生成结果为空（模型返回无有效 sections）"
                    + (f"；片段: {preview}…" if preview else "")
                )
            _validate_professional_sections(title, sections, slide_count=slide_count)
            return title, sections
        except Exception as exc:  # noqa: BLE001
            last_err = exc
            user = (
                f"{base_user}\n\n【上一版不合格，必须重写】\n"
                f"问题：{exc}\n"
                f"请严格输出 {slide_count} 节，与视频分镜逐页对应；书面语改写口播，"
                "叙事顺序不变；禁止互动引导和独立研报体例。"
            )
            print(f"[forum] 叙事长文第 {attempt + 1} 次不合格，准备重试：{exc}")
    raise RuntimeError(f"叙事长文生成多次不合格：{last_err}") from last_err


def _fallback_forum_sections(slides: list[dict], image_paths: list[Path]) -> list[dict]:
    sections: list[dict] = []
    for i, slide in enumerate(slides, start=1):
        h = (slide.get("headline") or "").strip()
        body = _expand_forum_section(slide)
        sections.append(
            {
                "headline": h,
                "body": body,
                "image": image_paths[i - 1] if i <= len(image_paths) else None,
            }
        )
    return sections


REPORT_FIGURE_SPECS: tuple[tuple[str, str], ...] = (
    (
        "图1：电子板块成交集中度",
        "Professional financial research chart on clean white background, 16:9 landscape. "
        "Left: rising bar chart labeled in Chinese '月度成交额' with trend '连续20个月超2万亿'. "
        "Right: donut chart with 28% slice highlighted labeled '占A股成交近28%'. "
        "Subtitle in small Chinese: '2026年5月 15.98万亿'. "
        "Corporate blue-gray palette, crisp vector infographic, Bloomberg-style data visualization. "
        "No hand-drawn doodles, no graph paper, no page numbers, no cartoon style.",
    ),
    (
        "图2：AI算力产业链传导",
        "Professional horizontal flow diagram, 16:9 landscape, clean white background. "
        "Left to right in Chinese: '大模型训练推理' -> '云端算力基础设施' -> "
        "'算力芯片/存储' -> 'PCB/先进封装' -> 'A股电子产业链'. "
        "Minimal blue arrows, flat vector business diagram, research report style. "
        "No hand-drawn style, no kitchen metaphors, no page numbers.",
    ),
    (
        "图3：景气与拥挤度矩阵",
        "Professional 2x2 quadrant matrix chart, 16:9 landscape, white background. "
        "X-axis Chinese '基本面兑现程度', Y-axis Chinese '交易拥挤度'. "
        "Highlight one dot in top-right quadrant labeled '电子板块'. "
        "Quadrant labels: high/low combinations in Chinese. "
        "Clean corporate infographic, blue accent, no doodles, no page numbers.",
    ),
)


def generate_report_figures(out_dir: Path, *, max_figures: int = 3) -> list[Path]:
    """为深度报告生成横版信息图（与短视频白板图分离）。"""
    from image_client import generate_image, save_b64_image

    load_env()
    images_dir = out_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    size = os.environ.get("AIVIDEO_REPORT_IMAGE_SIZE", "1536x1024")
    paths: list[Path] = []
    for i, (caption, prompt) in enumerate(REPORT_FIGURE_SPECS[:max_figures], start=1):
        print(f"[forum] 报告配图 {i}/{max_figures}：{caption}…")
        result = generate_image(prompt, size=size)
        dst = images_dir / f"{i:02d}.jpg"
        if result.get("b64_json"):
            save_b64_image(result["b64_json"], dst)
        elif result.get("url"):
            import urllib.request

            urllib.request.urlretrieve(result["url"], dst)
        if dst.is_file():
            paths.append(dst)
            print(f"[forum]   ✓ {dst.name} ({result.get('elapsed_s')}s)")
    return paths


def _parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if row and not all(set(c) <= {"-", ":"} for c in row):
            rows.append(row)
        i += 1
    return rows, i


def write_docx_from_post_md(post_md: Path, out_docx: Path, *, pack_dir: Path) -> Path:
    """把 post.md 转为 Word，表格与配图内嵌，便于复制发布。"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = post_md.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(11)

    i = 0
    fig_idx = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = _parse_md_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        table.rows[r].cells[c].text = cell
                doc.add_paragraph("")
            continue
        if stripped.startswith("**【插入配图"):
            m = re.search(r"`([^`]+)`", stripped)
            if m:
                rel = m.group(1)
                img_path = pack_dir / rel
                if img_path.is_file():
                    fig_idx += 1
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(6.2))
                    cap = doc.add_paragraph(f"图{fig_idx}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if stripped.startswith("---"):
            i += 1
            continue
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(("#", "|", "**【", "---")):
            para_lines.append(lines[i].strip())
            i += 1
        doc.add_paragraph("\n".join(para_lines))

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    return out_docx


def _render_forum_markdown(title: str, sections: list[dict]) -> str:
    lines = [f"# {title}", ""]
    has_risk_section = any(
        "风险" in str(sec.get("headline") or "") for sec in sections
    )
    for i, sec in enumerate(sections, start=1):
        h = str(sec.get("headline") or "").strip()
        body = str(sec.get("body") or "").strip()
        image = sec.get("image")
        if h:
            lines.append(f"## {h}")
            lines.append("")
        if body:
            lines.append(body)
            lines.append("")
        if image:
            lines.append(f"**【插入配图 {i}】** `images/{i:02d}.jpg`")
            lines.append("")
    if not has_risk_section and DISCLAIMER.strip():
        lines.extend(["---", "", DISCLAIMER, ""])
    return "\n".join(lines)


def _write_landscape_cover(src: Path, out_dir: Path) -> Path | None:
    """16:9 横封面（雪球首页推荐等），从竖封面居中偏上裁剪。"""
    if not src.is_file():
        return None
    try:
        from PIL import Image
    except ImportError:
        return None

    out_w = max(640, int(os.environ.get("AIVIDEO_FORUM_LANDSCAPE_W", "1280")))
    out_h = max(360, int(os.environ.get("AIVIDEO_FORUM_LANDSCAPE_H", "720")))
    try:
        focus_y = float(os.environ.get("AIVIDEO_FORUM_LANDSCAPE_FOCUS_Y", "0.38"))
    except ValueError:
        focus_y = 0.38
    focus_y = max(0.0, min(1.0, focus_y))

    img = Image.open(src).convert("RGB")
    w, h = img.size
    target_ratio = out_w / out_h
    if w / h >= target_ratio:
        new_h = h
        new_w = int(h * target_ratio)
        x0 = (w - new_w) // 2
        y0 = 0
    else:
        new_w = w
        new_h = int(w / target_ratio)
        x0 = 0
        y0 = int((h - new_h) * focus_y)
    crop = img.crop((x0, y0, x0 + new_w, y0 + new_h))
    if crop.size != (out_w, out_h):
        crop = crop.resize((out_w, out_h), Image.LANCZOS)

    dst = out_dir / "cover_landscape.jpg"
    crop.save(dst, "JPEG", quality=92)
    return dst if dst.is_file() else None


def build_forum_pack(
    script_path: Path,
    video_path: Path,
    out_dir: Path | None = None,
    *,
    allow_fallback: bool | None = None,
) -> dict:
    script, script_meta = _load_script_meta(script_path)
    title = _sanitize_forum_title((script.get("title") or "未命名").strip())
    slides = script.get("slides") or []
    out_dir = out_dir or forum_dir_for_video(video_path)
    out_dir.mkdir(parents=True, exist_ok=True)
    images_dir = out_dir / "images"

    image_paths = _copy_slide_images(script, images_dir, video_path, len(slides) or 4)
    cover_path, cover_src = _write_cover(script_path, video_path, out_dir)
    if not cover_path or not cover_path.is_file():
        for cand in image_paths:
            if cand.is_file():
                dst = out_dir / "cover.jpg"
                if _save_cover_jpg(cand, dst):
                    cover_path = dst
                    cover_src = cover_src or cand
                    print(f"[forum] cover.jpg 回退自配图: {cand.name}")
                    break
        if not cover_path or not cover_path.is_file():
            print(
                "[forum] ⚠️ 未生成 cover.jpg，知乎预填可能失败"
                "（请检查 ffmpeg 或 logs/.../cover.png）"
            )
    landscape_src = cover_src or cover_path
    landscape_path = (
        _write_landscape_cover(landscape_src, out_dir) if landscape_src else None
    )

    if os.environ.get("AIVIDEO_FORUM_PRO_ARTICLE", "1").strip().lower() in ("0", "false", "no"):
        sections = _fallback_forum_sections(slides, image_paths)
    else:
        try:
            title, sections = _generate_professional_forum_sections(
                script,
                slides=slides,
                image_paths=image_paths,
                article=script_meta.get("article") or script.get("article"),
                details=_details_for_forum(script_meta),
            )
        except Exception as exc:  # noqa: BLE001
            use_fallback = allow_fallback
            if use_fallback is None:
                use_fallback = os.environ.get("AIVIDEO_FORUM_ALLOW_FALLBACK", "0").strip().lower() in (
                    "1",
                    "true",
                    "yes",
                )
            if use_fallback:
                print(f"[forum] 叙事长文生成失败，回退口播扩展稿：{exc}")
                title = _sanitize_forum_title(str(script.get("title") or "未命名"))
                sections = _fallback_forum_sections(slides, image_paths)
            else:
                raise RuntimeError(
                    "叙事长文生成失败，已停止生成论坛图文，避免发布口播短稿。"
                    "如需临时回退旧稿，设置 AIVIDEO_FORUM_ALLOW_FALLBACK=1。"
                    f"原因：{exc}"
                ) from exc

    post_md = out_dir / "post.md"
    post_text = _render_forum_markdown(title, sections)
    post_md.write_text(post_text, encoding="utf-8")

    publish_sections = _build_readme_publish_sections(script)
    readme = f"""# 发布文案 · {video_path.name}

{publish_sections}"""
    (out_dir / "README.md").write_text(readme, encoding="utf-8")

    docx_path: Path | None = None
    if post_md.is_file():
        try:
            docx_path = write_docx_from_post_md(
                post_md, out_dir / "article.docx", pack_dir=out_dir,
            )
            print(f"[forum] Word 文稿：{docx_path}")
        except Exception as exc:  # noqa: BLE001
            print(f"[forum] Word 导出跳过: {exc}")

    return {
        "title": title,
        "out_dir": str(out_dir),
        "post_md": str(post_md),
        "article_docx": str(docx_path) if docx_path else "",
        "cover": str(cover_path) if cover_path else "",
        "cover_landscape": str(landscape_path) if landscape_path else "",
        "images": [str(p) for p in image_paths],
        "video": str(video_path),
    }


# 兼容旧引用
forum_out_dir_for_video = forum_dir_for_video
